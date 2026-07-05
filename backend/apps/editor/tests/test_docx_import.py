"""DOCX import fidelity: tables preserved, images materialised, scanned fallback.

Word uploads should behave like native Markdown notes — an editable body that
flows through the reader's TOC / typography / editing pipeline. These tests
pin the conversion (``convert_docx``) and the end-to-end import wiring.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from rest_framework.test import APIClient

from apps.editor.models import Attachment
from apps.editor.services.docx_import import convert_docx
from apps.knowledge.models import Document, KnowledgeBase
from apps.knowledge.serializers import detect_doc_format

User = get_user_model()

DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _png_bytes() -> bytes:
    from PIL import Image

    buf = BytesIO()
    Image.new("RGB", (12, 8), (200, 20, 20)).save(buf, format="PNG")
    return buf.getvalue()


def _make_docx(*, text: str = "Hello world", table: bool = False, image: bool = False) -> bytes:
    from docx import Document as Docx

    d = Docx()
    d.add_heading("Title One", level=1)
    d.add_heading("Section Two", level=2)
    if text:
        d.add_paragraph(text)
    if table:
        t = d.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "H1"
        t.cell(0, 1).text = "H2"
        t.cell(1, 0).text = "c1"
        t.cell(1, 1).text = "c2"
    if image:
        d.add_picture(BytesIO(_png_bytes()))
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# convert_docx unit behaviour
# --------------------------------------------------------------------------- #


def test_convert_preserves_atx_headings():
    md, imgs = convert_docx(_make_docx(text="body text"))
    assert "# Title One" in md
    assert "## Section Two" in md
    assert imgs == []


def test_convert_preserves_tables_as_wrapped_html():
    md, _ = convert_docx(_make_docx(table=True))
    assert '<div class="jz-table-wrap">' in md
    assert "<table" in md
    # cell text survives inside the raw table HTML
    assert "H1" in md and "c2" in md


def test_convert_extracts_embedded_image():
    md, imgs = convert_docx(_make_docx(image=True))
    assert len(imgs) == 1
    img = imgs[0]
    assert img.data  # real bytes
    assert img.token in md  # placeholder present as ![](token)
    assert img.ext in {".png", ".jpg"}


def test_convert_returns_empty_on_garbage():
    md, imgs = convert_docx(b"this is not a docx file")
    assert md == ""
    assert imgs == []


# --------------------------------------------------------------------------- #
# import endpoint end-to-end
# --------------------------------------------------------------------------- #


@pytest.fixture
def api_client():
    return APIClient()


@pytest.fixture
def owner():
    return User.objects.create_user("docxowner", "docx@e.com", "pass", is_staff=True)


@pytest.fixture
def kb(owner):
    return KnowledgeBase.objects.create(owner=owner, name="Docx KB", slug="docx-kb")


@pytest.mark.django_db
def test_import_docx_with_table_and_image(api_client, owner, kb, settings, tmp_path):
    settings.MEDIA_ROOT = str(tmp_path)
    api_client.force_authenticate(owner)
    f = SimpleUploadedFile(
        "report.docx", _make_docx(table=True, image=True), content_type=DOCX_CT
    )
    resp = api_client.post(
        reverse("api_v1:import-file"),
        data={"knowledge_base": kb.id, "file": f},
        format="multipart",
    )
    assert resp.status_code == 201, resp.content
    doc = Document.objects.get(pk=resp.data["id"])

    # Body-bearing → will flow through the Markdown reader (TOC/typography/edit).
    assert doc.raw_content.strip()
    assert doc.raw_content == doc.published_content  # import double-write
    assert '<div class="jz-table-wrap">' in doc.raw_content
    assert "# Title One" in doc.raw_content

    # The embedded image became a bound image Attachment and the body points at
    # /media/… (no dangling placeholder token).
    img_atts = doc.attachments.filter(kind=Attachment.KIND_IMAGE)
    assert img_atts.count() == 1
    assert "jzdocimg" not in doc.raw_content  # token was rewritten
    assert "/media/" in doc.raw_content

    # Original .docx retained + still detected as docx (badge unchanged).
    assert doc.attachments.filter(kind=Attachment.KIND_DOCUMENT).count() == 1
    assert detect_doc_format(doc) == "docx"


@pytest.mark.django_db
def test_import_scanned_docx_gets_fallback_body(
    api_client, owner, kb, settings, tmp_path, monkeypatch
):
    """A DOCX with no extractable text still yields a non-empty editable body."""
    settings.MEDIA_ROOT = str(tmp_path)
    # Force the "nothing extracted" path deterministically.
    monkeypatch.setattr(
        "apps.editor.services.docx_import.convert_docx", lambda blob: ("", [])
    )
    api_client.force_authenticate(owner)
    f = SimpleUploadedFile("scan.docx", _make_docx(text=""), content_type=DOCX_CT)
    resp = api_client.post(
        reverse("api_v1:import-file"),
        data={"knowledge_base": kb.id, "file": f},
        format="multipart",
    )
    assert resp.status_code == 201, resp.content
    doc = Document.objects.get(pk=resp.data["id"])
    assert doc.raw_content.strip()  # fallback body present
    assert "未能提取" in doc.raw_content
    # Original file kept so the reader can surface it at the bottom.
    assert doc.attachments.filter(kind=Attachment.KIND_DOCUMENT).count() == 1
    assert detect_doc_format(doc) == "docx"
