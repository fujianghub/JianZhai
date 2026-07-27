"""DOCX 导出大修回归：Heading 样式 / 图片 / 表格 / 插件降级 / TOC 域。"""
from __future__ import annotations

import pytest
from docx import Document as DocxDocument

from apps.exporter.scope import collect_for_scope
from apps.exporter.services import docx_export
from apps.exporter.tests.conftest import make_doc


def _export_docx(owner, kb, scope_kind="kb", target_id=None):
    scope = collect_for_scope(
        owner=owner, scope=scope_kind, target_id=target_id or kb.id
    )
    path, filename, mime = docx_export.export(scope)
    return DocxDocument(str(path)), filename


def _all_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@pytest.mark.django_db
def test_headings_use_real_heading_styles(owner, kb):
    make_doc(kb, "d", published="# 一级\n\n## 二级\n\n正文")
    document, _ = _export_docx(owner, kb)
    styles = [p.style.name for p in document.paragraphs]
    # 文档标题独占 Heading 1；正文标题降一级嵌套（镜像 PDF heading_shift）
    assert "Heading 1" in styles
    assert "Heading 2" in styles  # 正文 h1
    assert "Heading 3" in styles  # 正文 h2
    body_styles = [
        p.style.name for p in document.paragraphs if p.text.strip() in ("一级", "二级")
    ]
    assert body_styles == ["Heading 2", "Heading 3"]


@pytest.mark.django_db
def test_gfm_pipe_table_becomes_real_table(owner, kb):
    make_doc(kb, "d", published="| 姓名 | 得分 |\n| --- | --- |\n| 甲 | 90 |")
    document, _ = _export_docx(owner, kb)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "姓名"
    assert table.cell(1, 1).text == "90"


@pytest.mark.django_db
def test_html_color_table_extracted_not_dropped(owner, kb):
    html_table = (
        '<div class="jz-table-wrap"><table><tbody>'
        '<tr><td style="background:#fee">红格子</td><td>白格子</td></tr>'
        "</tbody></table></div>"
    )
    make_doc(kb, "d", published=f"前文\n\n{html_table}\n\n后文")
    document, _ = _export_docx(owner, kb)
    text = _all_text(document)
    assert "红格子" in text
    assert "白格子" in text


@pytest.mark.django_db
def test_image_embedded_from_media(owner, kb, settings, tmp_path):
    from PIL import Image

    settings.MEDIA_ROOT = str(tmp_path)
    (tmp_path / "uploads").mkdir()
    Image.new("RGB", (40, 20), "#cc0000").save(tmp_path / "uploads" / "p.png")
    make_doc(kb, "d", published="![示意图](/media/uploads/p.png)")
    document, _ = _export_docx(owner, kb)
    assert len(document.inline_shapes) == 1


@pytest.mark.django_db
def test_callout_and_strike_and_tasklist_degrade(owner, kb):
    src = ":::info 温馨提示\n盒内内容\n:::\n\n~~划掉~~\n\n- [x] 已办\n- [ ] 待办"
    make_doc(kb, "d", published=src)
    document, _ = _export_docx(owner, kb)
    text = _all_text(document)
    assert ":::" not in text  # callout 标记不再字面量泄漏
    assert "温馨提示" in text
    assert "盒内内容" in text
    assert "划掉" in text
    assert "☑" in text and "☐" in text
    assert "[x]" not in text


@pytest.mark.django_db
def test_layout_separators_degrade(owner, kb):
    src = ":::cols-2\n左列\n::col\n右列\n:::\n\n:::tabs\n::tab 方案A\n甲\n::tab 方案B\n乙\n:::"
    make_doc(kb, "d", published=src)
    document, _ = _export_docx(owner, kb)
    text = _all_text(document)
    assert "::col" not in text
    assert "::tab" not in text
    assert "左列" in text and "右列" in text
    assert "方案A" in text and "甲" in text and "乙" in text


@pytest.mark.django_db
def test_multi_doc_has_cover_and_toc_field(owner, kb):
    make_doc(kb, "a", published="one")
    make_doc(kb, "b", published="two")
    document, _ = _export_docx(owner, kb)
    styles = [p.style.name for p in document.paragraphs]
    assert "Title" in styles  # 封面标题
    assert "目录" in _all_text(document)
    xml = document.element.xml
    assert "fldSimple" in xml and "TOC" in xml  # Word 目录域


@pytest.mark.django_db
def test_html_format_doc_not_one_giant_paragraph(owner, kb):
    html_doc = "<html><body><h1>标题甲</h1><p>段一</p><p>段二</p></body></html>"
    make_doc(kb, "d", published=html_doc)
    document, _ = _export_docx(owner, kb)
    texts = [p.text for p in document.paragraphs if p.text.strip()]
    assert "段一" in texts and "段二" in texts  # 逐块成段，而非折叠为一行


@pytest.mark.django_db
def test_multi_doc_per_doc_toc_with_bookmarks(owner, kb):
    """多篇导出：每篇 meta 后有可点击「本篇目录」，标题挂书签、条目为内链。"""
    from apps.knowledge.models import Document

    make_doc(kb, "a", published="# 甲章\n\n## 甲节\n\n正文")
    make_doc(kb, "b", published="# 乙章\n\n正文")
    document, _ = _export_docx(owner, kb)
    doc_a = Document.objects.get(knowledge_base=kb, slug="a")
    xml = document.element.xml
    text = _all_text(document)
    assert text.count("本篇目录") == 2
    assert f'w:name="d{doc_a.id}_h0"' in xml  # 标题书签
    assert 'w:anchor="d' in xml  # 目录条目内链
    # 本篇目录出现在正文之前：目录条目「甲章」在标题段之前出现
    texts = [p.text for p in document.paragraphs]
    assert texts.index("本篇目录") < texts.index("甲章")


@pytest.mark.django_db
def test_single_doc_no_per_doc_toc(owner, kb):
    make_doc(kb, "only", published="# 甲\n\n正文")
    document, _ = _export_docx(owner, kb)
    assert "本篇目录" not in _all_text(document)


@pytest.mark.django_db
def test_html_doc_headings_get_styles_and_toc(owner, kb):
    """HTML 格式文档的 <h1-6> 转真 Heading 样式并进本篇目录（此前全退化纯文本）。"""
    html_doc = (
        "<html><body><h1>总纲</h1><p>引言</p>"
        "<h2>第一节</h2><p>内容一</p></body></html>"
    )
    make_doc(kb, "h", published=html_doc)
    make_doc(kb, "other", published="伴篇")
    document, _ = _export_docx(owner, kb)
    styled = {
        p.text.strip(): p.style.name
        for p in document.paragraphs
        if p.text.strip() in ("总纲", "第一节")
    }
    assert styled == {"总纲": "Heading 2", "第一节": "Heading 3"}
    assert "本篇目录" in _all_text(document)
    assert "引言" in _all_text(document)


@pytest.mark.django_db
def test_toc_placeholder_stripped(owner, kb):
    make_doc(kb, "d", published="[TOC]\n\n# 一\n\n正文")
    document, _ = _export_docx(owner, kb)
    assert "[TOC]" not in _all_text(document)


@pytest.mark.django_db
def test_east_asia_fonts_configured(owner, kb):
    make_doc(kb, "d", published="正文")
    document, _ = _export_docx(owner, kb)
    normal = document.styles["Normal"]
    rfonts = normal.element.rPr.rFonts
    assert rfonts.get(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
    ) == "宋体"
