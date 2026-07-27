"""DOCX export — walk the markdown-it token stream and emit python-docx elements.

Headings use real ``Heading N`` styles (Word 导航窗格 + 自动目录都依赖它们)；
多篇合订带封面 + TOC 域（Word 打开后更新域即生成带页码目录）。图片内嵌、
GFM 管道表转真表格、原生 HTML 表格提取文字矩阵、callout 降级为标签段。
仍为已知限制：无 OMML 公式（LaTeX 源码 run）、表格颜色/间距不保留。
"""
from __future__ import annotations

import html as html_lib
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from apps.knowledge.serializers import detect_doc_format
from markdown_it import MarkdownIt
from mdit_py_plugins.container import container_plugin
from mdit_py_plugins.gfm import gfm_plugin

from ..scope import ExportScope
from . import card_placeholders, common
from .markdown_preprocess import (
    _HTML_COMMENT,
    apply_yuque_compat_mode,
    map_outside_fenced_code_blocks,
    normalize_latex_delimiters,
    recover_yuque_diagram_comments,
    unglue_container_fences,
)
from .markdown_render import install_math_rules

_md = MarkdownIt("commonmark", {"breaks": True})
gfm_plugin(_md)  # tables + strikethrough + autolink + tasklists（checked 落在 list_item_open.meta）
# Callout containers parse into container tokens (label paragraph + normal
# inner content) instead of leaking ``:::info`` markers as literal text.
container_plugin(
    _md,
    "callout",
    marker=":",
    validate=lambda params, name: bool(params.strip()),
)
# 数学 tokenizer（escape 后、emphasis 前拦截 ``$``）——没有它，公式里的
# ``_``/``*``/``\`` 会被 CommonMark 吃掉，docx 输出残破字面量。docx 不产
# OMML（已知限制），公式以 LaTeX 源码 run 原样保留，可复制回编辑器。
install_math_rules(_md)

_LAYOUT_COL_SEP = re.compile(r"^::col\s*$", re.M)
_LAYOUT_TAB_SEP = re.compile(r"^::tab[ \t]*(.*)$", re.M)


def _degrade_layout_separators(src: str) -> str:
    """docx 无分栏/标签页概念：``::col`` 行删除（列内容顺排），``::tab 标签``
    降级为加粗标签行。避免分隔符以字面量泄漏进 Word 正文。"""
    out = _LAYOUT_COL_SEP.sub("", src)
    return _LAYOUT_TAB_SEP.sub(lambda m: f"**{m.group(1).strip()}**" if m.group(1).strip() else "", out)


def _docx_preprocess(body: str) -> str:
    """docx 专用预处理：与 preprocess_markdown 同源，但布局块不转 HTML
    （docx 的 html_block 只能降级，转 HTML 反而会丢内容）。"""
    out = recover_yuque_diagram_comments(body or "")
    out = _HTML_COMMENT.sub("", out)
    out = map_outside_fenced_code_blocks(out, unglue_container_fences)
    out = map_outside_fenced_code_blocks(out, normalize_latex_delimiters)
    out = map_outside_fenced_code_blocks(out, apply_yuque_compat_mode)
    out = map_outside_fenced_code_blocks(out, _degrade_layout_separators)
    return out


def export(scope: ExportScope) -> tuple[Path, str, str]:
    docx = DocxDocument()
    _setup_styles(docx)

    # 卡片占位符降级为普通链接行（docx 本就丢 href，标题文字得以保留）
    card_titles = card_placeholders.doc_titles_for(
        common.doc_export_body(d) for d in scope.documents
    )

    multi = len(scope.documents) > 1
    if multi:
        _emit_front_matter(docx, scope)

    for idx, doc in enumerate(scope.documents):
        if idx > 0 or multi:
            docx.add_page_break()
        title_para = docx.add_paragraph(style="Heading 1")
        title_para.add_run(doc.title)
        meta_para = docx.add_paragraph()
        meta_run = meta_para.add_run(
            f"{doc.knowledge_base.name}"
            + (f" · {doc.published_at:%Y-%m-%d}" if doc.published_at else "")
        )
        meta_run.italic = True
        meta_run.font.size = Pt(10)

        body = common.doc_export_body(doc)
        if detect_doc_format(doc) == "html":
            # 逐块断行提取，而非整篇折叠成一个巨型段落
            for line in _html_to_lines(body):
                docx.add_paragraph(line)
        else:
            body = card_placeholders.degrade_card_placeholders(
                body, doc_titles=card_titles
            )
            tokens = _md.parse(_docx_preprocess(body))
            _render_tokens(docx, tokens)

    path = common.reserve_export_path(".docx")
    docx.save(path)
    return (
        path,
        common.build_export_filename(scope, ".docx"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _emit_front_matter(docx: DocxDocument, scope: ExportScope) -> None:
    """封面（大类 / 标题 / 篇数）+ TOC 域。Word 打开后「更新域」即得带页码目录
    （字段依赖真 Heading 样式——正是本导出器切换到 Heading N 的原因）。"""
    from django.utils import timezone

    category = scope.kb.category
    if category is not None:
        cat_para = docx.add_paragraph()
        cat_run = cat_para.add_run(category.name)
        cat_run.font.size = Pt(12)
        cat_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para = docx.add_paragraph(style="Title")
    title_para.add_run(scope.label)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_para = docx.add_paragraph()
    meta_run = meta_para.add_run(
        f"共 {len(scope.documents)} 篇 · {timezone.localtime():%Y-%m-%d} · 简斋"
    )
    meta_run.font.size = Pt(10)
    meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    docx.add_page_break()
    toc_label = docx.add_paragraph()
    toc_label_run = toc_label.add_run("目录")
    toc_label_run.bold = True
    toc_label_run.font.size = Pt(16)
    toc_para = docx.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = "（在 Word 中选中此行按 F9 更新域，即可生成带页码目录）"
    run_el.append(text_el)
    fld.append(run_el)
    toc_para._p.append(fld)


def _setup_styles(docx: DocxDocument) -> None:
    """Normal + Heading + Title 显式设 eastAsia 字体——默认西文样式会让
    中文回退成宋体杂拌；标题用雅黑、正文用宋体的常规中文排版。"""
    _style_fonts(docx, "Normal", "Calibri", "宋体", size=Pt(11))
    for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4",
                 "Heading 5", "Heading 6", "Title"):
        try:
            _style_fonts(docx, name, "Calibri Light", "微软雅黑")
        except KeyError:
            continue


def _style_fonts(docx, style_name: str, ascii_name: str, east_asia: str, size=None):
    style = docx.styles[style_name]
    style.font.name = ascii_name
    if size is not None:
        style.font.size = size
    # ``style.font.name`` only sets the ascii/hAnsi slots; CJK glyphs read the
    # eastAsia slot which python-docx has no high-level API for.
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def _render_tokens(docx: DocxDocument, tokens) -> None:
    """Lightweight pass over markdown-it tokens. Tracks lists/blockquote/heading state."""
    list_stack: list[tuple[str, int]] = []  # (kind: 'bullet'|'ordered', counter)
    quote_depth = 0
    task_prefix: str | None = None  # gfm tasklists: checked 态挂在 list_item_open.meta

    i = 0
    while i < len(tokens):
        t = tokens[i]
        tt = t.type

        if tt == "list_item_open" and t.meta.get("checked") is not None:
            task_prefix = "☑ " if t.meta["checked"] else "☐ "
            i += 1
            continue

        if tt == "heading_open":
            level = min(int(t.tag[1]), 6)  # h1..h6
            inline = tokens[i + 1]
            try:
                para = docx.add_paragraph(style=f"Heading {level}")
            except KeyError:
                para = docx.add_paragraph()
            _emit_inline_runs(para, inline)
            i += 3  # heading_open, inline, heading_close
            continue

        if tt == "paragraph_open":
            inline = tokens[i + 1]
            para = docx.add_paragraph()
            if task_prefix:
                para.add_run(task_prefix)
                task_prefix = None
            _emit_inline_runs(para, inline)
            if list_stack:
                kind, _ = list_stack[-1]
                depth = len(list_stack)
                base = "List Bullet" if kind == "bullet" else "List Number"
                name = base if depth == 1 else f"{base} {min(depth, 3)}"
                try:
                    para.style = docx.styles[name]
                except KeyError:
                    para.style = docx.styles[base]
                    para.paragraph_format.left_indent = Pt(12 * (depth - 1))
            elif quote_depth:
                try:
                    para.style = docx.styles["Quote"]
                except KeyError:
                    para.paragraph_format.left_indent = Pt(12 * quote_depth)
            i += 3
            continue

        if tt == "bullet_list_open":
            list_stack.append(("bullet", 0))
            i += 1
            continue
        if tt == "ordered_list_open":
            list_stack.append(("ordered", 0))
            i += 1
            continue
        if tt in ("bullet_list_close", "ordered_list_close"):
            list_stack.pop()
            i += 1
            continue

        if tt == "blockquote_open":
            quote_depth += 1
            i += 1
            continue
        if tt == "blockquote_close":
            quote_depth = max(0, quote_depth - 1)
            i += 1
            continue

        if tt.startswith("container_callout_open"):
            label = (t.info or "").strip()
            if label:
                para = docx.add_paragraph()
                run = para.add_run(f"▍{label}")
                run.bold = True
            i += 1
            continue
        if tt.startswith("container_callout_close"):
            i += 1
            continue

        if tt == "code_block" or tt == "fence":
            para = docx.add_paragraph()
            run = para.add_run(t.content.rstrip("\n"))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            para.paragraph_format.left_indent = Pt(12)
            i += 1
            continue

        if tt == "hr":
            para = docx.add_paragraph("─" * 40)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        if tt == "math_block":
            para = docx.add_paragraph()
            run = para.add_run(f"$${t.content}$$")
            run.font.name = "Cambria Math"
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue

        if tt == "table_open":
            rows, i = _extract_md_table(tokens, i)
            _emit_table(docx, rows)
            continue

        if tt == "html_block":
            _emit_html_block(docx, t.content)
            i += 1
            continue

        # list_item_open / close are no-ops; their inner paragraph_open handles formatting
        i += 1


def _extract_md_table(tokens, i: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    row: list[str] | None = None
    depth = 0
    j = i
    while j < len(tokens):
        tk = tokens[j]
        if tk.type == "table_open":
            depth += 1
        elif tk.type == "table_close":
            depth -= 1
            if depth == 0:
                j += 1
                break
        elif tk.type == "tr_open":
            row = []
        elif tk.type == "tr_close":
            if row is not None:
                rows.append(row)
            row = None
        elif tk.type in ("th_open", "td_open") and row is not None:
            inline = tokens[j + 1] if j + 1 < len(tokens) else None
            row.append(
                _inline_text(inline) if inline is not None and inline.type == "inline" else ""
            )
        j += 1
    return rows, j


def _emit_table(docx: DocxDocument, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    if not cols:
        return
    table = docx.add_table(rows=len(rows), cols=cols)
    try:
        table.style = docx.styles["Table Grid"]
    except KeyError:
        pass
    for r_i, r in enumerate(rows):
        for c_i in range(cols):
            cell = table.cell(r_i, c_i)
            cell.text = r[c_i] if c_i < len(r) else ""
            if r_i == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


class _HtmlTableExtractor(HTMLParser):
    """原生 HTML 表格（带色表的条件序列化形态）→ 纯文字矩阵。样式丢（已知
    限制），但内容不再整块蒸发。"""

    def __init__(self):
        super().__init__()
        self.rows: list[list[str]] = []
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def handle_starttag(self, tag, attrs):
        if tag == "tr":
            self._row = []
        elif tag in ("td", "th"):
            self._cell = []

    def handle_endtag(self, tag):
        if tag in ("td", "th") and self._cell is not None and self._row is not None:
            self._row.append(re.sub(r"\s+", " ", "".join(self._cell)).strip())
            self._cell = None
        elif tag == "tr" and self._row is not None:
            self.rows.append(self._row)
            self._row = None

    def handle_data(self, data):
        if self._cell is not None:
            self._cell.append(data)


_HTML_IMG_SRC = re.compile(r"<img\b[^>]*?src=[\"'](?P<url>/media/[^\"']+)[\"']", re.I)
_BLOCK_BREAK = re.compile(r"(?i)<(?:/p|/div|/h[1-6]|/li|/tr|/table|br\s*/?)\s*>")


def _emit_html_block(docx: DocxDocument, content: str) -> None:
    """html_block 三分支：表格 → 文字矩阵真表格；图片 → 内嵌；其余 → 剥
    标签按行落段。此前 html_block 无分支直接跳过——带色表格整块蒸发。"""
    if "<table" in content.lower():
        parser = _HtmlTableExtractor()
        try:
            parser.feed(content)
        except Exception:
            parser.rows = []
        if parser.rows:
            _emit_table(docx, parser.rows)
            return
    imgs = list(_HTML_IMG_SRC.finditer(content))
    if imgs:
        for m in imgs:
            _add_picture(docx.add_paragraph(), m.group("url"), alt="")
        return
    for line in _html_to_lines(content):
        docx.add_paragraph(line)


def _html_to_lines(html_src: str) -> list[str]:
    """块级闭合标签断行 → 剥标签 → 反转义。修掉「HTML 文档折叠成单个巨型
    段落」的旧行为（\\s+ 折叠丢光了全部换行）。"""
    txt = _BLOCK_BREAK.sub("\n", html_src or "")
    txt = re.sub(r"<[^>]+>", " ", txt)
    txt = html_lib.unescape(txt)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in txt.split("\n")]
    return [line for line in lines if line]


def _add_picture(para, url: str, alt: str) -> None:
    path = common._resolve_media_path(url)
    if path is None:
        if alt or url:
            para.add_run(f"[图片：{alt or url}]")
        return
    run = para.add_run()
    try:
        pic = run.add_picture(str(path))
    except Exception:
        para.add_run(f"[图片无法嵌入：{alt or path.name}]")
        return
    max_w = Inches(6)
    if pic.width > max_w:
        ratio = max_w / pic.width
        pic.height = int(pic.height * ratio)
        pic.width = max_w


def _emit_inline_runs(para, inline_token) -> None:
    """Walk an inline token's children, emitting runs with bold/italic/code marks."""
    bold = italic = strike = False
    for child in inline_token.children or []:
        ct = child.type
        if ct == "strong_open":
            bold = True
        elif ct == "strong_close":
            bold = False
        elif ct in ("em_open", "emphasis_open"):
            italic = True
        elif ct in ("em_close", "emphasis_close"):
            italic = False
        elif ct == "s_open":
            strike = True
        elif ct == "s_close":
            strike = False
        elif ct == "code_inline":
            run = para.add_run(child.content)
            run.font.name = "Consolas"
            continue
        elif ct == "math_inline":
            run = para.add_run(f"${child.content}$")
            run.font.name = "Cambria Math"
            continue
        elif ct == "image":
            _add_picture(
                para,
                child.attrGet("src") or "",
                "".join(c.content for c in (child.children or []) if c.type == "text"),
            )
        elif ct == "html_inline":
            lowered = child.content.lower()
            if "<input" in lowered and "checkbox" in lowered:
                para.add_run("☑ " if "checked" in lowered else "☐ ")
            # 其余行内标签（<font>/<span> 等）静默剥掉，文字子节点照常输出
        elif ct == "softbreak":
            para.add_run(" ")
        elif ct == "hardbreak":
            para.add_run("\n")
        elif ct == "text":
            run = para.add_run(child.content)
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
        elif ct == "link_open":
            # Render the link text inline; href is dropped (Word footnotes are out of scope here).
            pass
        elif ct == "link_close":
            pass


def _inline_text(inline_token) -> str:
    return "".join(
        c.content
        for c in (inline_token.children or [])
        if c.type in ("text", "code_inline", "math_inline")
    )
